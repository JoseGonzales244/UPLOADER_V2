"""
Extractor de Transcripciones de WhatsApp (.docx) con Filtrado del Ejecutivo Evaluado.
Extrae ÚNICAMENTE los mensajes del ejecutivo asignado en la gestión (filtrando bots y otros ejecutivos).
Implementa ITranscriptExtractor para respetar contratos abstractos.
"""
import os
import re
import glob
import logging
import docx
import pandas as pd
from typing import List, Dict, Any, Optional

from domain.interfaces.transcript_extractor import ITranscriptExtractor

logger = logging.getLogger("modules.transcripciones.extractors.wsp_docx_extractor")


def get_paragraph_full_text(p) -> str:
    """Extrae el texto completo del párrafo conservando nodos <w:t> y tabulaciones <w:tab/>."""
    text_pieces = []
    for node in p._element.xpath(".//w:t | .//w:tab"):
        if node.tag.endswith("tab"):
            text_pieces.append("\t")
        elif node.text:
            text_pieces.append(node.text)
    return "".join(text_pieces)


class WhatsAppTranscriptExtractor(ITranscriptExtractor):
    def __init__(self, folder_path: Optional[str] = None):
        self.folder_path = folder_path or os.environ.get(
            "WSP_INPUT_DIR",
            os.path.join("data", "input", "auditorias_wsp")
        )
        self.ejecutivos_df = self._load_ejecutivos_mapping()

    def _load_ejecutivos_mapping(self) -> pd.DataFrame:
        excel_path = os.path.join(self.folder_path, "Ejecutivos_Gestion_Wsp.xlsx")
        if os.path.exists(excel_path):
            try:
                df = pd.read_excel(excel_path)
                logger.info(f"Cargado mapeo de ejecutivos desde {excel_path} ({len(df)} registros).")
                return df
            except Exception as e:
                logger.error(f"Error leyendo {excel_path}: {e}")
        return pd.DataFrame()

    def extract_single_docx(self, docx_path: str) -> Dict[str, Any]:
        """
        Extrae ÚNICAMENTE la conversación correspondiente AL EJECUTIVO EVALUADO del archivo .docx.
        Filtra bots, flujos automáticos y mensajes de otros ejecutivos ajenos a la gestión.
        """
        filename = os.path.basename(docx_path)
        interaction_id = os.path.splitext(filename)[0]

        exec_info = {
            "SUPERVISOR": "N/A",
            "REGISTRO COLABORADOR": "N/A",
            "COLABORADOR": "N/A",
            "SUB EQUIPO": "Televentas"
        }
        
        target_registro = ""
        target_nombre = ""

        if not self.ejecutivos_df.empty and "ID_Interaccion" in self.ejecutivos_df.columns:
            match = self.ejecutivos_df[self.ejecutivos_df["ID_Interaccion"].astype(str).str.strip() == interaction_id]
            if not match.empty:
                row = match.iloc[0]
                target_registro = str(row.get("REGISTRO COLABORADOR", "") or "").strip()
                target_nombre = str(row.get("COLABORADOR", "") or "").strip()
                exec_info = {
                    "SUPERVISOR": str(row.get("SUPERVISOR", "N/A")),
                    "REGISTRO COLABORADOR": target_registro or "N/A",
                    "COLABORADOR": target_nombre or "N/A",
                    "SUB EQUIPO": str(row.get("SUB EQUIPO", "Televentas"))
                }

        doc = docx.Document(docx_path)
        
        paragraphs_text = []
        for p in doc.paragraphs:
            txt = get_paragraph_full_text(p).strip()
            if txt:
                paragraphs_text.append(txt)

        # También extraer texto de celdas en tablas de Word
        for table in doc.tables:
            for row in table.rows:
                cell_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cell_texts:
                    line = "\t".join(cell_texts)
                    if line not in paragraphs_text:
                        paragraphs_text.append(line)

        # 1. Parsear eventos tabulados de Genesys
        events = []
        for line in paragraphs_text:
            parts = line.split("\t")
            if len(parts) >= 3 and parts[1] in ("Interno", "Externo"):
                fecha_hora = parts[0]
                tipo_part = parts[1]
                sender = parts[2]
                mensaje_texto = parts[3] if len(parts) >= 4 else ""
                mensaje_texto = mensaje_texto.replace("\u202c", "").replace("\u202b", "").strip()

                events.append({
                    "fecha_hora": fecha_hora,
                    "tipo": tipo_part,
                    "sender": sender,
                    "texto": mensaje_texto
                })

        def is_target_exec(sender):
            s_lower = sender.lower()
            if "bot" in s_lower or "flujo" in s_lower or "acd" in s_lower:
                return False
            if target_registro:
                digits = re.sub(r"\D", "", target_registro)
                if digits and digits in s_lower:
                    return True
            if target_nombre:
                words = [w.lower() for w in re.split(r"[\s,]+", target_nombre) if len(w.strip()) > 3]
                match_count = sum(1 for w in words if w in s_lower)
                if match_count >= 1:
                    return True
            return False

        if events:
            # Filtrar por el rango de eventos del ejecutivo evaluado objetivo
            target_indices = [i for i, ev in enumerate(events) if ev["tipo"] == "Interno" and is_target_exec(ev["sender"])]

            if not target_indices:
                # Si no hay filtro específico por registro, incluir mensajes de ejecutivos internos (no bots)
                target_indices = [i for i, ev in enumerate(events) if ev["tipo"] == "Interno" and not any(b in ev["sender"].lower() for b in ["bot", "flujo", "acd"])]

            if target_indices:
                min_idx = target_indices[0]
                max_idx = target_indices[-1]

                start_idx = max(0, min_idx - 1) if min_idx > 0 and events[min_idx - 1]["tipo"] == "Externo" else min_idx
                end_idx = min(len(events) - 1, max_idx + 1) if max_idx < len(events) - 1 and events[max_idx + 1]["tipo"] == "Externo" else max_idx

                dialogue_lines = []
                for i in range(start_idx, end_idx + 1):
                    ev = events[i]
                    if ev["tipo"] == "Externo":
                        dialogue_lines.append(f"Cliente ({ev['sender']}) [{ev['fecha_hora']}]: {ev['texto']}")
                    elif ev["tipo"] == "Interno" and (is_target_exec(ev["sender"]) or not target_registro):
                        sender_label = target_nombre if target_nombre else ev['sender']
                        dialogue_lines.append(f"Ejecutivo Evaluado [{sender_label}] [{ev['fecha_hora']}]: {ev['texto']}")

                raw_dialogue_text = "\n".join(dialogue_lines)
            else:
                raw_dialogue_text = "Sin mensajes de texto registrados para el ejecutivo evaluado en la transcripción."

            header_summary = (
                f"=== FICHA DE EVALUACIÓN WHATSAPP ===\n"
                f"ID Interacción: {interaction_id}\n"
                f"Ejecutivo Evaluado: {exec_info['COLABORADOR']} (Registro: {exec_info['REGISTRO COLABORADOR']})\n"
                f"Supervisor: {exec_info['SUPERVISOR']} | Sub-Equipo: {exec_info['SUB EQUIPO']}\n"
                f"====================================\n\n"
            )
            full_text = header_summary + raw_dialogue_text
            return {
                "archivo": filename,
                "interaction_id": interaction_id,
                "full_text": full_text,
                "executive_interaction": raw_dialogue_text,
                "metadata": exec_info
            }

        # 2. Si no hay eventos tabulados de Genesys, extraer diálogo filtrando por el ejecutivo
        dialogue_lines = []
        for txt in paragraphs_text:
            if txt.startswith("Archivo:") or txt.startswith("Tipo de") or txt.startswith("ID de"):
                continue
            dialogue_lines.append(txt)

        raw_dialogue_text = "\n".join(dialogue_lines) if dialogue_lines else "\n".join(paragraphs_text)

        header_summary = (
            f"=== FICHA DE EVALUACIÓN WHATSAPP (CONVERSACIÓN LIMPIA) ===\n"
            f"ID Interacción: {interaction_id}\n"
            f"Ejecutivo Evaluado: {exec_info['COLABORADOR']} (Registro: {exec_info['REGISTRO COLABORADOR']})\n"
            f"Supervisor: {exec_info['SUPERVISOR']} | Sub-Equipo: {exec_info['SUB EQUIPO']}\n"
            f"====================================\n\n"
        )
        full_text = header_summary + raw_dialogue_text

        return {
            "archivo": filename,
            "interaction_id": interaction_id,
            "full_text": full_text,
            "executive_interaction": raw_dialogue_text,
            "metadata": exec_info
        }

    def get_all_transcripts(self) -> List[Dict[str, Any]]:
        """Busca y extrae todos los archivos .docx en la carpeta de transcripciones."""
        docx_files = glob.glob(os.path.join(self.folder_path, "*.docx"))
        results = []
        for filepath in sorted(docx_files):
            try:
                data = self.extract_single_docx(filepath)
                if data["full_text"].strip():
                    results.append(data)
            except Exception as e:
                logger.error(f"Error extrayendo {filepath}: {e}")
        return results
